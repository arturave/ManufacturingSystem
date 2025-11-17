#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Advanced Report Generator with Thumbnails Support
Generates Excel and PDF reports with embedded product thumbnails
"""

import os
import io
import datetime
from typing import Dict, List, Any, Optional
from pathlib import Path
from PIL import Image

# Excel imports
import openpyxl
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.drawing.image import Image as OpenpyxlImage
from openpyxl.utils import get_column_letter
from openpyxl.chart import PieChart, BarChart, Reference

# Word/PDF imports
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
try:
    from docx2pdf import convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False
    print("Warning: docx2pdf not available. PDF generation will be disabled.")

# Settings
from settings_manager import get_settings_manager

class AdvancedReportGenerator:
    """Advanced report generator with thumbnail support"""

    def __init__(self):
        """Initialize report generator"""
        self.settings_manager = get_settings_manager()
        self.settings = self.settings_manager.settings

    def generate_excel_report(
        self,
        data: List[Dict[str, Any]],
        output_path: str,
        title: str = "Raport",
        customer_info: Optional[Dict[str, Any]] = None,
        order_info: Optional[Dict[str, Any]] = None
    ) -> bool:
        """
        Generate Excel report with thumbnails

        Args:
            data: List of product/part data
            output_path: Path for output file
            title: Report title
            customer_info: Customer information
            order_info: Order information

        Returns:
            bool: True if successful
        """
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Raport główny"

            # Company logo and header
            row = 1
            if self.settings.report_include_thumbnails:
                logo_path = self.settings_manager.get_active_logo_path()
                if logo_path and os.path.exists(logo_path):
                    try:
                        img = OpenpyxlImage(logo_path)
                        img.width = 150
                        img.height = 60
                        ws.add_image(img, 'A1')
                        row = 5
                    except Exception as e:
                        print(f"Could not add logo: {e}")

            # Title
            ws.merge_cells(f'A{row}:H{row}')
            title_cell = ws[f'A{row}']
            title_cell.value = title
            title_cell.font = Font(bold=True, size=16)
            title_cell.alignment = Alignment(horizontal="center")
            row += 2

            # Customer info
            if customer_info:
                ws[f'A{row}'] = "Klient:"
                ws[f'A{row}'].font = Font(bold=True)
                ws[f'B{row}'] = customer_info.get('name', '')
                row += 1

                if customer_info.get('address'):
                    ws[f'A{row}'] = "Adres:"
                    ws[f'A{row}'].font = Font(bold=True)
                    ws[f'B{row}'] = customer_info.get('address', '')
                    row += 1

                if customer_info.get('nip'):
                    ws[f'A{row}'] = "NIP:"
                    ws[f'A{row}'].font = Font(bold=True)
                    ws[f'B{row}'] = customer_info.get('nip', '')
                    row += 1

                row += 1

            # Order info
            if order_info:
                ws[f'A{row}'] = "Numer zamówienia:"
                ws[f'A{row}'].font = Font(bold=True)
                ws[f'B{row}'] = order_info.get('order_number', '')
                row += 1

                ws[f'A{row}'] = "Data:"
                ws[f'A{row}'].font = Font(bold=True)
                ws[f'B{row}'] = order_info.get('date', datetime.datetime.now().strftime('%Y-%m-%d'))
                row += 2

            # Table headers
            headers = ['Lp.', 'Miniatura', 'Nazwa produktu', 'Indeks', 'Materiał',
                      'Grubość [mm]', 'Ilość', 'Cena jedn. [PLN]', 'Wartość [PLN]']

            # Remove Miniatura column if thumbnails are disabled
            if not self.settings.report_include_thumbnails:
                headers.remove('Miniatura')

            for col_idx, header in enumerate(headers, 1):
                cell = ws.cell(row=row, column=col_idx, value=header)
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
                cell.alignment = Alignment(horizontal="center", vertical="center")
                cell.border = Border(
                    left=Side(style='thin'),
                    right=Side(style='thin'),
                    top=Side(style='thin'),
                    bottom=Side(style='thin')
                )

            row += 1
            data_start_row = row

            # Add data rows
            total_value = 0
            for idx, item in enumerate(data, 1):
                col = 1

                # Lp.
                ws.cell(row=row, column=col, value=idx)
                col += 1

                # Thumbnail (if enabled)
                if self.settings.report_include_thumbnails:
                    # Leave empty cell for thumbnail
                    col += 1

                    # Add thumbnail if available
                    if item.get('thumbnail_data') or item.get('thumb_data'):
                        try:
                            thumb_data = item.get('thumbnail_data') or item.get('thumb_data')
                            if thumb_data:
                                img = OpenpyxlImage(io.BytesIO(thumb_data))
                                img.width = 60
                                img.height = 40
                                ws.add_image(img, f'{get_column_letter(2)}{row}')
                                ws.row_dimensions[row].height = 45
                        except Exception as e:
                            print(f"Failed to add thumbnail: {e}")

                # Product name
                ws.cell(row=row, column=col, value=item.get('name', ''))
                col += 1

                # Index
                ws.cell(row=row, column=col, value=item.get('idx_code', ''))
                col += 1

                # Material
                material = ''
                if item.get('materials_dict'):
                    material = item['materials_dict'].get('name', '')
                elif item.get('material'):
                    material = item['material']
                ws.cell(row=row, column=col, value=material)
                col += 1

                # Thickness
                thickness = item.get('thickness_mm', 0)
                ws.cell(row=row, column=col, value=thickness)
                col += 1

                # Quantity
                qty = item.get('qty', 1)
                ws.cell(row=row, column=col, value=qty)
                col += 1

                # Unit price
                unit_price = item.get('price', 0) or item.get('cost_per_unit', 0)
                cell = ws.cell(row=row, column=col, value=unit_price)
                cell.number_format = '#,##0.00'
                col += 1

                # Total value
                total_item = unit_price * qty
                total_value += total_item
                cell = ws.cell(row=row, column=col, value=total_item)
                cell.number_format = '#,##0.00'

                # Apply borders
                for c in range(1, col + 1):
                    ws.cell(row=row, column=c).border = Border(
                        left=Side(style='thin'),
                        right=Side(style='thin'),
                        top=Side(style='thin'),
                        bottom=Side(style='thin')
                    )

                row += 1

            # Total row
            ws.merge_cells(f'A{row}:{get_column_letter(len(headers)-2)}{row}')
            total_label_cell = ws[f'A{row}']
            total_label_cell.value = "RAZEM:"
            total_label_cell.font = Font(bold=True)
            total_label_cell.alignment = Alignment(horizontal="right")

            total_value_cell = ws[f'{get_column_letter(len(headers))}{row}']
            total_value_cell.value = total_value
            total_value_cell.font = Font(bold=True)
            total_value_cell.number_format = '#,##0.00'
            total_value_cell.border = Border(
                top=Side(style='double'),
                bottom=Side(style='double')
            )

            # Auto-fit columns
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter

                # Special widths for specific columns
                if column_letter == 'B' and self.settings.report_include_thumbnails:
                    ws.column_dimensions[column_letter].width = 12
                elif column_letter == 'C':
                    ws.column_dimensions[column_letter].width = 30
                else:
                    for cell in column:
                        try:
                            if cell.value:
                                max_length = max(max_length, len(str(cell.value)))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 40)
                    ws.column_dimensions[column_letter].width = adjusted_width

            # Save workbook
            wb.save(output_path)
            return True

        except Exception as e:
            print(f"Error generating Excel report: {e}")
            return False

    def generate_word_report(
        self,
        data: List[Dict[str, Any]],
        output_path: str,
        title: str = "Raport",
        customer_info: Optional[Dict[str, Any]] = None,
        order_info: Optional[Dict[str, Any]] = None
    ) -> bool:
        """
        Generate Word report with thumbnails

        Args:
            data: List of product/part data
            output_path: Path for output file
            title: Report title
            customer_info: Customer information
            order_info: Order information

        Returns:
            bool: True if successful
        """
        try:
            doc = Document()

            # Set margins
            sections = doc.sections
            for section in sections:
                section.left_margin = Cm(2)
                section.right_margin = Cm(2)
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)

            # Add logo
            if self.settings.report_include_thumbnails:
                logo_path = self.settings_manager.get_active_logo_path()
                if logo_path and os.path.exists(logo_path):
                    try:
                        doc.add_picture(logo_path, width=Inches(2.5))
                    except Exception as e:
                        print(f"Could not add logo: {e}")

            # Title
            title_para = doc.add_heading(title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Customer info
            if customer_info:
                doc.add_heading("Dane klienta", level=2)
                doc.add_paragraph(f"Nazwa: {customer_info.get('name', '')}")
                if customer_info.get('address'):
                    doc.add_paragraph(f"Adres: {customer_info.get('address', '')}")
                if customer_info.get('nip'):
                    doc.add_paragraph(f"NIP: {customer_info.get('nip', '')}")
                doc.add_paragraph()

            # Order info
            if order_info:
                doc.add_heading("Dane zamówienia", level=2)
                doc.add_paragraph(f"Numer: {order_info.get('order_number', '')}")
                doc.add_paragraph(f"Data: {order_info.get('date', datetime.datetime.now().strftime('%Y-%m-%d'))}")
                doc.add_paragraph()

            # Products table
            doc.add_heading("Produkty", level=2)

            # Determine number of columns based on thumbnail setting
            if self.settings.report_include_thumbnails:
                table = doc.add_table(rows=1, cols=7)
                headers = ['Lp.', 'Miniatura', 'Nazwa', 'Indeks', 'Materiał', 'Ilość', 'Cena']
            else:
                table = doc.add_table(rows=1, cols=6)
                headers = ['Lp.', 'Nazwa', 'Indeks', 'Materiał', 'Ilość', 'Cena']

            table.style = 'Table Grid'

            # Set headers
            hdr_cells = table.rows[0].cells
            for idx, header in enumerate(headers):
                hdr_cells[idx].text = header
                # Style header
                for paragraph in hdr_cells[idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)

            # Add data rows
            total_value = 0
            for idx, item in enumerate(data, 1):
                row_cells = table.add_row().cells
                col = 0

                # Lp.
                row_cells[col].text = str(idx)
                col += 1

                # Thumbnail (if enabled)
                if self.settings.report_include_thumbnails:
                    if item.get('thumbnail_data') or item.get('thumb_data'):
                        try:
                            thumb_data = item.get('thumbnail_data') or item.get('thumb_data')
                            if thumb_data:
                                paragraph = row_cells[col].paragraphs[0]
                                run = paragraph.add_run()
                                run.add_picture(io.BytesIO(thumb_data), width=Inches(0.8))
                        except Exception as e:
                            print(f"Failed to add thumbnail: {e}")
                    col += 1

                # Name
                row_cells[col].text = item.get('name', '')
                col += 1

                # Index
                row_cells[col].text = item.get('idx_code', '')
                col += 1

                # Material
                material = ''
                if item.get('materials_dict'):
                    material = item['materials_dict'].get('name', '')
                elif item.get('material'):
                    material = item['material']
                row_cells[col].text = material
                col += 1

                # Quantity
                row_cells[col].text = str(item.get('qty', 1))
                col += 1

                # Price
                price = item.get('price', 0) or item.get('cost_per_unit', 0)
                row_cells[col].text = f"{price:.2f} PLN"
                total_value += price * item.get('qty', 1)

            # Add total
            doc.add_paragraph()
            total_para = doc.add_paragraph()
            total_para.add_run(f"Wartość całkowita: {total_value:.2f} PLN").bold = True
            total_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Save document
            doc.save(output_path)
            return True

        except Exception as e:
            print(f"Error generating Word report: {e}")
            return False

    def generate_pdf_report(
        self,
        data: List[Dict[str, Any]],
        output_path: str,
        title: str = "Raport",
        customer_info: Optional[Dict[str, Any]] = None,
        order_info: Optional[Dict[str, Any]] = None
    ) -> bool:
        """
        Generate PDF report by converting from Word

        Args:
            data: List of product/part data
            output_path: Path for output PDF file
            title: Report title
            customer_info: Customer information
            order_info: Order information

        Returns:
            bool: True if successful
        """
        if not DOCX2PDF_AVAILABLE:
            print("docx2pdf module not available. Cannot generate PDF.")
            return False

        try:
            # Generate Word document first
            temp_docx = output_path.replace('.pdf', '_temp.docx')
            if self.generate_word_report(data, temp_docx, title, customer_info, order_info):
                # Convert to PDF
                convert(temp_docx, output_path)

                # Remove temporary Word file
                try:
                    os.remove(temp_docx)
                except:
                    pass

                return True
            return False

        except Exception as e:
            print(f"Error generating PDF report: {e}")
            return False

    def generate_comprehensive_report(
        self,
        data: List[Dict[str, Any]],
        output_dir: str,
        base_name: str,
        formats: List[str] = ['xlsx', 'pdf'],
        **kwargs
    ) -> Dict[str, bool]:
        """
        Generate reports in multiple formats

        Args:
            data: List of product/part data
            output_dir: Directory for output files
            base_name: Base name for files
            formats: List of formats to generate ('xlsx', 'docx', 'pdf')
            **kwargs: Additional arguments for reports

        Returns:
            Dict with format as key and success status as value
        """
        results = {}

        # Ensure output directory exists
        os.makedirs(output_dir, exist_ok=True)

        # Generate each format
        for format_type in formats:
            output_path = os.path.join(output_dir, f"{base_name}.{format_type}")

            if format_type == 'xlsx':
                results['xlsx'] = self.generate_excel_report(data, output_path, **kwargs)
            elif format_type == 'docx':
                results['docx'] = self.generate_word_report(data, output_path, **kwargs)
            elif format_type == 'pdf':
                results['pdf'] = self.generate_pdf_report(data, output_path, **kwargs)
            else:
                results[format_type] = False

        return results